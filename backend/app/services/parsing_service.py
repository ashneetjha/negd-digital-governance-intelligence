"""
Parsing service extracts raw text with section headings and page numbers
from PDF and DOCX files.
"""
from __future__ import annotations

import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from app.utils.logger import get_logger

logger = get_logger(__name__)


class ParsedPage:
    def __init__(self, text: str, page_number: int, section_heading: str | None = None):
        self.text = text
        self.page_number = page_number
        self.section_heading = section_heading


# =========================================================
# PDF PARSER
# =========================================================

def extract_from_pdf(file_path: str) -> list[ParsedPage]:
    pages: list[ParsedPage] = []
    current_heading: str | None = None

    try:
        import fitz
    except ImportError:
        fitz = None

    try:
        if fitz:
            doc = fitz.open(file_path)

            for page_num, page in enumerate(doc, start=1):
                blocks = page.get_text("dict").get("blocks", [])
                text_parts = []

                for block in blocks:
                    if block.get("type") != 0:
                        continue
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            text = (span.get("text") or "").strip()
                            if not text:
                                continue

                            if span.get("size", 0) >= 13 or "Bold" in span.get("font", ""):
                                current_heading = text

                            text_parts.append(text)

                full_text = " ".join(text_parts).strip()

                if full_text:
                    pages.append(
                        ParsedPage(
                            text=full_text,
                            page_number=page_num,
                            section_heading=current_heading,
                        )
                    )

            doc.close()

            if not pages:
                raise ValueError("PDF parsing failed")

            logger.info("PDF parsed", file=file_path, pages=len(pages))
            return pages

        # fallback pypdf
        from pypdf import PdfReader

        reader = PdfReader(file_path)

        for page_num, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue

            lines = [l.strip() for l in text.splitlines() if l.strip()]

            if lines:
                head = lines[0]
                if len(head) < 120:
                    current_heading = head

            pages.append(
                ParsedPage(
                    text=" ".join(lines),
                    page_number=page_num,
                    section_heading=current_heading,
                )
            )

        if not pages:
            raise ValueError("PDF parsing failed")

        return pages

    except Exception as e:
        logger.error("PDF parse error", error=str(e))
        raise


# =========================================================
# DOCX RAW FALLBACK
# =========================================================

def _docx_raw_text_fallback(file_path: str) -> list[str]:
    try:
        with zipfile.ZipFile(file_path, "r") as zf:
            xml = zf.read("word/document.xml")

        root = ET.fromstring(xml)
        raw = " ".join(root.itertext())
        raw = re.sub(r"\s+", " ", raw)

        sentences = re.split(r"(?<=[.!?])\s+", raw)
        result = [s.strip() for s in sentences if s.strip()]
        if not result:
            raise ValueError("DOCX raw fallback returned empty text")
        return result

    except Exception as e:
        logger.warning("DOCX fallback failed", error=str(e))
        return []


# =========================================================
# DOCX PARSER
# =========================================================

def parse_docx(file_path: str) -> list[str]:
    from docx import Document

    doc = Document(file_path)
    blocks: list[str] = []

    # paragraphs
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            blocks.append(txt)

    # tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                c.text.strip() for c in row.cells if c.text.strip()
            )
            if row_text:
                blocks.append(row_text)

    logger.info("DOCX parsed", blocks=len(blocks))

    # fallback
    if len(blocks) < 3:
        fallback = _docx_raw_text_fallback(file_path)
        if fallback:
            blocks = fallback

    if not blocks:
        raise ValueError("DOCX parsing failed")

    return blocks


_DOCX_HEADING_RE = re.compile(r"^[A-Z][A-Za-z0-9\s\-:]{3,100}$")


def extract_from_docx(file_path: str) -> list[ParsedPage]:
    try:
        blocks = parse_docx(file_path)

        pages: list[ParsedPage] = []
        current_heading: str | None = None

        for i, block in enumerate(blocks, start=1):
            txt = block.strip()

            if " | " not in txt and len(txt) < 120 and _DOCX_HEADING_RE.match(txt):
                current_heading = txt

            pages.append(
                ParsedPage(
                    text=txt,
                    page_number=i,
                    section_heading=current_heading,
                )
            )

        logger.info("DOCX pages created", count=len(pages))
        return pages

    except Exception as e:
        logger.error("DOCX parse error", error=str(e))
        raise


# =========================================================
# MAIN ROUTER
# =========================================================

def parse_document(file_path: str) -> list[ParsedPage]:
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        pages = extract_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        pages = extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file: {ext}")

    if not pages:
        raise ValueError("Parsing failed - empty output")

    return pages